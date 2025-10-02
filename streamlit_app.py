import io
import re
import string
from datetime import date, timedelta

import streamlit as st
import pandas as pd
import pdfplumber
from docx import Document

st.set_page_config(page_title="Review Pasien — Streamlit", page_icon="🦷", layout="wide")

# ===== DPJP Canon =====
DPJP_CANON = [
    "Dr. drg. Andi Tajrin, M.Kes., Sp.B.M.M., Subsp. C.O.M.(K)",
    "drg. Mukhtar Nur Anam Sp.B.M.M.",
    "drg. Husnul Basyar, Sp. B.M.M.",
    "drg. Abul Fauzi, Sp.B.M.M., Subsp.T.M.T.M.J.(K)",
    "drg. M. Irfan Rasul, Ph.D., Sp.B.M.M., Subsp.C.O.M.(K)",
    "drg. Mohammad Gazali, MARS., Sp.B.M.M., Subsp.T.M.T.M.J.(K)",
    "drg. Timurwati, Sp.B.M.M.",
    "drg. Husni Mubarak, Sp. B.M.M.",
    "drg. Nurwahida, M.K.G., Sp.B.M.M., Subsp.C.O.M(K)",
    "drg. Hadira, M.K.G., Sp.B.M.M., Subsp.C.O.M(K)",
    "drg. Carolina Stevanie, Sp.B.M.M.",
    "drg. Yossy Yoanita Ariestiana, M.KG., Sp.B.M.M., Subsp.Ortognat-D (K)",
]

# ===== Fuzzy DPJP =====
def _norm_doctor(s: str) -> str:
    if not s:
        return ""
    # perbaiki typo umum & variasi gelar
    s = s.replace("drg..", "drg.")
    s = s.replace("Sp. BM", "Sp.BM").replace("Sp. BMM", "Sp.BMM")
    s = s.replace("Sp.BM(K)", "Sp.BMM").replace("Sp.BMM(K)", "Sp.BMM")
    s = s.replace("Sp.BM", "Sp.BMM")
    # normalisasi huruf & token
    s = s.upper()
    # KUNCI PERBAIKAN: non-huruf jadi SPASI (bukan dihapus)
    import re
    s = re.sub(r"[^A-Z]+", " ", s)
    return " ".join(s.split())
    
def _token_jaccard(a: str, b: str) -> float:
    ta = set(_norm_doctor(a).split())
    tb = set(_norm_doctor(b).split())
    if not ta or not tb: return 0.0
    return len(ta & tb) / len(ta | tb)

def map_doctor_to_canonical(raw: str, candidates=DPJP_CANON, threshold: float = 0.35) -> str:
    best, best_score = "", 0.0
    for c in candidates:
        sc = _token_jaccard(raw, c)
        if sc > best_score:
            best, best_score = c, sc
    return best if best_score >= threshold else ""

# ===== Helpers =====
ID_MONTHS = {
    "JANUARI": 1, "FEBRUARI": 2, "MARET": 3, "APRIL": 4, "MEI": 5, "JUNI": 6,
    "JULI": 7, "AGUSTUS": 8, "SEPTEMBER": 9, "OKTOBER": 10, "NOVEMBER": 11, "DESEMBER": 12
}
ROMAN = {"I":1,"V":5,"X":10,"L":50,"C":100}

def roman_to_int(s: str) -> int:
    s = re.sub(r"[^IVXLC]", "", s.upper())
    if not s: return 0
    total = 0
    prev = 0
    for ch in reversed(s):
        val = ROMAN.get(ch, 0)
        if val < prev: total -= val
        else:
            total += val; prev = val
    return total

def fmt_rm(rm: str) -> str:
    digits = re.sub(r"\D", "", rm or "")
    digits = digits.zfill(6)[:6]
    return f"{digits[0:2]}.{digits[2:4]}.{digits[4:6]}"

def extract_period_date_from_text(text: str):
    m = re.search(r"PERIODE\s+(\d{1,2})\s+([A-Z]+)\s+(\d{4})", text.upper())
    if not m: return None
    d = int(m.group(1)); mon_name = m.group(2).strip(); y = int(m.group(3))
    mon = ID_MONTHS.get(mon_name)
    if not mon: return None
    try:
        return date(y, mon, d)
    except Exception:
        return None

def replace_gigi(text: str, gigi: str | None) -> str:
    if not (gigi and str(gigi).strip()):
        return text
    return re.sub(r"(?i)(\bgigi\s*)xx\b", r"\1" + str(gigi).strip(), text)

def compute_kontrol_text(kontrol_tpl: str, diagnosa_text: str, base_date):
    if not base_date: return kontrol_tpl
    mk = re.search(r"\bPOD\s+([IVXLC]+)\b", kontrol_tpl, flags=re.IGNORECASE)
    md = re.search(r"\bPOD\s+([IVXLC]+)\b", diagnosa_text or "", flags=re.IGNORECASE)
    if not mk: return kontrol_tpl

    pod_k = roman_to_int(mk.group(1))
    pod_d = roman_to_int(md.group(1)) if md else 0

    offset = pod_k - pod_d
    if offset < 0: offset = 0
    target = base_date + timedelta(days=offset)

    if pod_k == 3 and target.weekday() == 6:  # Sunday
        pod_k = 4
        target = target + timedelta(days=1)
        kontrol_tpl = re.sub(r"\bPOD\s+[IVXLC]+\b", "POD IV", kontrol_tpl, flags=re.IGNORECASE)

    date_str = target.strftime("%d/%m/%Y")
    if re.search(r"\([^)]*\)", kontrol_tpl):
        return re.sub(r"\([^)]*\)", f"({date_str})", kontrol_tpl)
    else:
        return f"{kontrol_tpl} ({date_str})"

# ===== Templates =====
B = "•⁠  ⁠"
LABELS = {
    "nama": "Nama            : ",
    "tgl":  f"{B}Tanggal lahir  : ",
    "rm":   f"{B}RM                   : ",
    "diag": f"{B}Diagnosa        : ",
    "tind": f"{B}Tindakan        : ",
    "kont": f"{B}Kontrol           : ",
    "dpjp": f"{B}DPJP               : ",
    "telp": f"{B}No. Telp.         : ",
    "opr":  f"{B}Operator         : ",
}
VISIT_TEMPLATES = {
    "(Pilih)": dict(diagnosa="", tindakan=[], kontrol=""),
    "Kunjungan 1": dict(
        diagnosa="",
        tindakan=["Konsultasi", "Periapikal X-ray gigi xx / OPG X-Ray"],
        kontrol="Pro ekstraksi gigi xx dalam lokal anestesi / Pro odontektomi gigi xx dalam lokal anestesi (xx/04/2025)",
    ),
    "Kunjungan 2": dict(
        diagnosa="Impaksi gigi xx kelas xx posisi xx Mesioangular / Gangren pulpa gigi xx / Gangren radiks gigi xx",
        tindakan=[
            "Odontektomi gigi xx dalam lokal anestesi",
            "ekstraksi gigi xx dengan penyulit dalam lokal anestesi",
            "ekstraksi gigi xx dengan open methode dalam lokal anestesi",
        ],
        kontrol="POD III (xx/04/2025)",
    ),
    "Kunjungan 3": dict(
        diagnosa="POD III Ekstraksi gigi xx dalam lokal anestesi / POD III Odontektomi gigi xx dalam lokal anestesi",
        tindakan=["Cuci luka intraoral dengan NaCl 0,9%"],
        kontrol="POD VII (xx/04/2025)",
    ),
    "Kunjungan 4": dict(
        diagnosa="POD VII Odontektomi gigi xx dalam lokal anestesi / POD VII Ekstraksi gigi xx dalam lokal anestesi",
        tindakan=["Cuci luka intra oral dengan NaCl 0,9%", "Aff hecting"],
        kontrol="POD XIV (xx/04/2025)",
    ),
    "Kunjungan 5": dict(
        diagnosa="POD XIV Ekstraksi gigi xx dalam lokal anestesi / POD XIV Odontektomi gigi xx dalam lokal anestesi",
        tindakan=["Kontrol luka post operasi", "Rujuk balik FKTP"],
        kontrol="-",
    ),
}

def normalize_visit(text: str) -> str:
    t = (text or "").strip()
    if not t: return "(Pilih)"
    if t.isdigit() and t in {"1","2","3","4","5"}:
        return f"Kunjungan {t}"
    low = t.lower()
    for k in VISIT_TEMPLATES.keys():
        if low == k.lower():
            return k
    return t

def build_block(no, row, visit_key, base_date):
    tpl_key = normalize_visit(visit_key or row.get("visit") or "(Pilih)")
    tpl = VISIT_TEMPLATES.get(tpl_key, VISIT_TEMPLATES["(Pilih)"])
    diagnosa = tpl["diagnosa"]
    tindakan = list(tpl["tindakan"])
    kontrol  = tpl["kontrol"]

    gigi = (row.get("gigi") or "").strip()
    diagnosa = replace_gigi(diagnosa, gigi)
    tindakan = [replace_gigi(t, gigi) for t in tindakan]

    kontrol = compute_kontrol_text(kontrol, diagnosa, base_date)

    dpjp_full = (row.get("DPJP (auto)") or "").strip()
    telp = (row.get("telp") or "").strip()
    operator = (row.get("operator") or "").strip()

    L = LABELS
    lines = []
    lines.append(f"{no}. {L['nama']}{row['Nama']}")
    lines.append(f"{L['tgl']}{row['Tgl Lahir']}")
    lines.append(f"{L['rm']}{fmt_rm(row['No. RM'])}")
    lines.append(f"{L['diag']}{diagnosa}")

    if tpl_key == "Kunjungan 3" and len(tindakan) == 1:
        lines.append(f"{L['tind']}{tindakan[0]}")
    else:
        lines.append(f"{L['tind']}")
        for t in tindakan:
            lines.append(f"    * {t}")

    lines.append(f"{L['kont']}{kontrol}")
    lines.append(f"{L['dpjp']}{dpjp_full}")
    lines.append(f"{L['telp']}{telp}")
    lines.append(f"{L['opr']}{operator}")
    return "\n".join(lines)

# ===== PDF Parser =====
def parse_pdf_to_rows_and_period_bytes(pdf_bytes: bytes):
    rows = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        full_text = ""
        for p in pdf.pages:
            txt = p.extract_text() or ""
            full_text += txt + "\n"

    period_date = extract_period_date_from_text(full_text)

    pat = re.compile(
        r"(?P<rm>\d{5,6})\s+"
        r"(?P<nopen>\d{8,18})\s+"
        r"(?P<name>.+?)\s+"
        r"(?P<sex>[LP])(?=\s+[0-3]\d-\d{2}-\d{4})\s+"
        r"(?P<dob>[0-3]\d-\d{2}-\d{4})",
        re.DOTALL
    )
    matches = list(pat.finditer(full_text))

    for i, m in enumerate(matches):
        start = m.start()
        end   = matches[i+1].start() if i+1 < len(matches) else len(full_text)
        block = full_text[start:end]

        raw_name = m.group("name")
        name_clean = re.sub(r"[ \t\r\f\v]+", " ", raw_name).replace("\n", " ")
        name_clean = re.sub(r"\s{2,}", " ", name_clean).strip().title()

        doc_m = re.search(r"(drg[^\n]+)", block, flags=re.IGNORECASE)
        dokter_raw = doc_m.group(1).strip() if doc_m else ""
        dokter_raw = re.split(r"\s\d{2}:\d{2}:\d{2}|ANJUNGAN|KLINIK|BELUM|,00|\.00|,0",
                              dokter_raw, 1, flags=re.IGNORECASE)[0].strip()
        dokter_raw = dokter_raw.rstrip(",;")
        dpjp_auto = map_doctor_to_canonical(dokter_raw)

        rows.append({
            "No. RM": m.group("rm"),
            "Nama":   name_clean,
            "Tgl Lahir": m.group("dob").replace("-", "/"),
            "DPJP (auto)": dpjp_auto,
            "checked": False,
            "visit": "(Pilih)",
            "gigi": "",
            "telp": "",
            "operator": "",
        })

    uniq = {}
    for r in rows:
        key = (r["No. RM"], r["Nama"], r["Tgl Lahir"])
        if key not in uniq:
            uniq[key] = r
    final = list(uniq.values())

    out = []
    for i, r in enumerate(final, start=1):
        rr = dict(r); rr["No."] = i; out.append(rr)

    return out, period_date

# ===== UI =====
st.title("🦷 Review Pasien (Ported) — Streamlit")
st.caption("Porting dari app Tkinter: parsing PDF → tabel → template WA")

with st.expander("Catatan & aturan format", expanded=False):
    st.markdown(
        "- **Nama multi-baris utuh**, NOPEN 8–18 digit, RM → `XX.XX.XX`.\n"
        "- **PERIODE** dipakai sebagai base tanggal kontrol.\n"
        "- Kunjungan 3: **Tindakan** satu baris (tanpa bullet).\n"
        "- Kontrol otomatis dari **POD** (dengan aturan Minggu → POD IV / +1 hari).\n"
    )

uploaded = st.file_uploader("Upload PDF laporan", type=["pdf"])

@st.cache_data(show_spinner=False)
def _parse_cached(pdf_bytes: bytes):
    return parse_pdf_to_rows_and_period_bytes(pdf_bytes)

if uploaded is not None:
    data = uploaded.read()
    try:
        rows, period_date = _parse_cached(data)
    except Exception as e:
        st.error(f"Gagal membaca PDF: {e}")
        st.stop()

    if not rows:
        st.error("PDF tidak terbaca / pola tidak cocok.")
        st.stop()

    per_str = period_date.strftime("%d/%m/%Y") if period_date else "—"
    st.success(f"Ditemukan {len(rows)} pasien — PERIODE: **{per_str}** — file: **{uploaded.name}**")

    df = pd.DataFrame(rows, columns=["No.","Nama","Tgl Lahir","No. RM","DPJP (auto)","visit","gigi","telp","operator","checked"])
    st.markdown("### Tabel pasien (editable)")
    edited = st.data_editor(
        df,
        column_config={
            "checked": st.column_config.CheckboxColumn("✓"),
            "visit": st.column_config.TextColumn("Kunjungan"),
            "gigi": st.column_config.TextColumn("Gigi"),
            "telp": st.column_config.TextColumn("Telp"),
            "operator": st.column_config.TextColumn("Operator"),
        },
        hide_index=True,
        use_container_width=True,
        num_rows="fixed",
        height=380,
    )

    # Build preview blocks
    sel = edited[edited["checked"] == True].copy()
    preview_blocks = []
    if not sel.empty:
        for _, r in sel.sort_values("No.").iterrows():
            rdict = {
                "Nama": r["Nama"],
                "Tgl Lahir": r["Tgl Lahir"],
                "No. RM": r["No. RM"],
                "DPJP (auto)": r["DPJP (auto)"],
                "visit": r["visit"],
                "gigi": r["gigi"],
                "telp": r["telp"],
                "operator": r["operator"],
            }
            block = build_block(int(r["No."]), rdict, r["visit"], period_date)
            preview_blocks.append(block)

    col1, col2 = st.columns([3,2], gap="large")
    with col1:
        st.markdown("### Preview teks final")
        if preview_blocks:
            final_text = "\n\n".join(preview_blocks)
            st.text_area("Teks hasil", final_text, height=360)
        else:
            final_text = ""
            st.info("Centang baris pasien yang ingin digenerate.")

        if final_text:
            # Download TXT
            st.download_button(
                "⬇️ Download TXT",
                data=final_text.encode("utf-8"),
                file_name="laporan_pasien.txt",
                mime="text/plain",
                use_container_width=True
            )

            # Download DOCX
            buf = io.BytesIO()
            doc = Document()
            for part in final_text.split("\n\n"):
                if part.strip():
                    for line in part.splitlines():
                        doc.add_paragraph(line)
                    doc.add_paragraph("")
            doc.save(buf)
            st.download_button(
                "⬇️ Download DOCX",
                data=buf.getvalue(),
                file_name="laporan_pasien.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

    with col2:
        st.markdown("### Opsi")
        st.write("- Ubah nilai **Kunjungan/Gigi/Telp/Operator** langsung di tabel.")
        st.write("- Kolom **✓** untuk memilih pasien yang akan digenerate.")
        st.write("- DPJP terisi otomatis via *fuzzy mapping* dari PDF.")

st.divider()
st.caption("Made for Streamlit Cloud — pdfplumber + python-docx")
